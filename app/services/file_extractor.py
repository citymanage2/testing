"""
Extract text content from uploaded files for passing to Claude as text.
Supports: DOCX, XLSX, XML (Grand CAD), TIFF (convert to PNG).
PDF and images are handled directly by the Anthropic API.
"""
import io
import base64


def extract_docx_text(data: bytes) -> str:
    """Extract plain text from a DOCX file."""
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_xlsx_text(data: bytes) -> str:
    """Extract content from an XLSX file as tab-separated text."""
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Лист: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            # Skip fully empty rows
            if any(c.strip() for c in cells):
                parts.append("\t".join(cells))
    wb.close()
    return "\n".join(parts)


def extract_xml_text(data: bytes, file_name: str) -> str:
    """Extract content from XML (Grand CAD / EDC exports)."""
    try:
        from lxml import etree
        root = etree.fromstring(data)
        # Collect all text nodes
        texts = []
        for elem in root.iter():
            tag = etree.QName(elem.tag).localname if "{" in elem.tag else elem.tag
            if elem.text and elem.text.strip():
                texts.append(f"{tag}: {elem.text.strip()}")
        return "\n".join(texts)
    except Exception:
        # Fallback: return raw XML as text (truncated if huge)
        text = data.decode("utf-8", errors="replace")
        if len(text) > 200_000:
            text = text[:200_000] + "\n... [файл обрезан]"
        return text


def convert_tiff_to_png_base64(data: bytes) -> str:
    """Convert TIFF image to PNG and return as base64 string."""
    from PIL import Image
    img = Image.open(io.BytesIO(data))
    out = io.BytesIO()
    img.save(out, format="PNG")
    return base64.standard_b64encode(out.getvalue()).decode()


def file_to_claude_part(file_name: str, mime_type: str, data: bytes) -> dict:
    """
    Convert an uploaded file into the appropriate Claude message content block.
    Returns a dict ready to be inserted into messages[].content[].
    """
    # ── Images ──────────────────────────────────────────────────────────────
    if mime_type in ("image/jpeg", "image/png", "image/gif", "image/webp"):
        return {
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": mime_type,
                "data": base64.standard_b64encode(data).decode(),
            },
        }

    if mime_type == "image/tiff":
        png_b64 = convert_tiff_to_png_base64(data)
        return {
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": "image/png",
                "data": png_b64,
            },
        }

    # ── PDF ─────────────────────────────────────────────────────────────────
    if mime_type == "application/pdf":
        return {
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": base64.standard_b64encode(data).decode(),
            },
        }

    # ── DOCX ────────────────────────────────────────────────────────────────
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        try:
            text = extract_docx_text(data)
        except Exception as e:
            text = f"[Не удалось извлечь текст из {file_name}: {e}]"
        return {"type": "text", "text": f"=== Содержимое файла {file_name} ===\n{text}"}

    # ── XLSX ────────────────────────────────────────────────────────────────
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        try:
            text = extract_xlsx_text(data)
        except Exception as e:
            text = f"[Не удалось извлечь текст из {file_name}: {e}]"
        return {"type": "text", "text": f"=== Содержимое файла {file_name} ===\n{text}"}

    # ── XML ─────────────────────────────────────────────────────────────────
    if mime_type in ("text/xml", "application/xml"):
        try:
            text = extract_xml_text(data, file_name)
        except Exception as e:
            text = f"[Не удалось извлечь текст из {file_name}: {e}]"
        return {"type": "text", "text": f"=== Содержимое файла {file_name} ===\n{text}"}

    # ── Fallback ─────────────────────────────────────────────────────────────
    return {"type": "text", "text": f"[Файл: {file_name}, тип: {mime_type} — содержимое недоступно]"}
